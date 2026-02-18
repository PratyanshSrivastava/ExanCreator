from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import copy


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = 'w:' + edge
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_run_font(run, font_name='Times New Roman', size_pt=12, bold=False, italic=False, underline=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def add_para(doc, text='', bold=False, italic=False, underline=False,
             font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, font_name='Times New Roman', line_spacing=None):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if text:
        run = para.add_run(text)
        set_run_font(run, font_name, font_size, bold, italic, underline)
    return para


def add_run_to_para(para, text, bold=False, italic=False, underline=False,
                    font_size=12, font_name='Times New Roman'):
    run = para.add_run(text)
    set_run_font(run, font_name, font_size, bold, italic, underline)
    return run


def add_horizontal_line(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def set_page_margins(doc, top=1, bottom=1, left=1, right=1):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def add_header_footer(doc, metadata):
    section = doc.sections[0]
    
    # Footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    set_info = metadata.get('set', '')
    if set_info:
        run_set = footer_para.add_run(f'{set_info}   |   ')
        set_run_font(run_set, size_pt=9)
    
    run_pg = footer_para.add_run('Page ')
    set_run_font(run_pg, size_pt=9)
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_pg._r.append(fldChar_begin)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run_pg._r.append(instrText)
    
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_pg._r.append(fldChar_end)
    
    run_of = footer_para.add_run(' of ')
    set_run_font(run_of, size_pt=9)
    
    fldChar_begin2 = OxmlElement('w:fldChar')
    fldChar_begin2.set(qn('w:fldCharType'), 'begin')
    run_of._r.append(fldChar_begin2)
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    run_of._r.append(instrText2)
    
    fldChar_end2 = OxmlElement('w:fldChar')
    fldChar_end2.set(qn('w:fldCharType'), 'end')
    run_of._r.append(fldChar_end2)


def build_header_table(doc, metadata):
    """Build the 3-column header table with school name, exam info."""
    school_name = metadata.get('schoolName', '').upper()
    school_addr1 = metadata.get('schoolAddressLine1', '')
    school_addr2 = metadata.get('schoolAddressLine2', '')
    exam_type = metadata.get('examType', '')
    academic_year = metadata.get('academicYear', '')
    class_name = metadata.get('class', '')
    subject = metadata.get('subject', '')
    subject_code = metadata.get('subjectCode', '')
    set_val = metadata.get('set', '')
    max_marks = metadata.get('maxMarks', '')
    duration = metadata.get('durationMinutes', '')
    
    # Duration display
    if duration:
        try:
            d = int(duration)
            if d % 60 == 0:
                dur_str = f"{d//60} Hr{'s' if d//60 > 1 else ''}"
            elif d % 60 == 30:
                dur_str = f"{d//60}.5 Hrs"
            else:
                dur_str = f"{d} Min"
        except:
            dur_str = str(duration)
    else:
        dur_str = ''

    # Main header table (school info)
    header_table = doc.add_table(rows=3, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.style = 'Table Grid'
    
    # Row 0: School name and exam
    cell = header_table.rows[0].cells[0]
    cell.width = Inches(6.3)
    # Remove all borders
    set_cell_border(cell, top={'val': 'none'}, bottom={'val': 'none'}, 
                    left={'val': 'none'}, right={'val': 'none'})
    
    if school_name:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(school_name)
        set_run_font(r, size_pt=16, bold=True)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
    
    if exam_type:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(exam_type.upper())
        set_run_font(r, size_pt=14, bold=True)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
    
    if school_addr1 or school_addr2:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr = ', '.join(filter(None, [school_addr1, school_addr2]))
        r = p.add_run(addr)
        set_run_font(r, size_pt=10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
    
    # Row 1: Class, Subject, Set, Marks, Time
    cell1 = header_table.rows[1].cells[0]
    set_cell_border(cell1, top={'val': 'single', 'sz': 6}, bottom={'val': 'single', 'sz': 6},
                    left={'val': 'none'}, right={'val': 'none'})
    
    p = cell1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    
    info_parts = []
    if class_name:
        info_parts.append(f'Class: {class_name}')
    if subject:
        sub_str = subject
        if subject_code:
            sub_str += f' (Code: {subject_code})'
        info_parts.append(f'Subject: {sub_str}')
    if set_val:
        info_parts.append(f'Set: {set_val}')
    
    info_line1 = '     |     '.join(info_parts)
    r = p.add_run(info_line1)
    set_run_font(r, size_pt=11, bold=True)
    
    p2 = cell1.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    
    info_parts2 = []
    if max_marks:
        info_parts2.append(f'Max. Marks: {max_marks}')
    if dur_str:
        info_parts2.append(f'Time: {dur_str}')
    if academic_year:
        info_parts2.append(f'Academic Year: {academic_year}')
    
    info_line2 = '     |     '.join(info_parts2)
    r2 = p2.add_run(info_line2)
    set_run_font(r2, size_pt=11, bold=True)
    
    # Row 2: Instructions
    cell2 = header_table.rows[2].cells[0]
    set_cell_border(cell2, top={'val': 'single', 'sz': 6}, bottom={'val': 'single', 'sz': 6},
                    left={'val': 'none'}, right={'val': 'none'})
    
    instructions = doc._body.tables[0].rows[2].cells[0] if False else cell2
    return header_table, cell2


def add_instructions(cell, instructions_list):
    if not instructions_list:
        return
    p = cell.add_paragraph()
    r = p.add_run('GENERAL INSTRUCTIONS:')
    set_run_font(r, size_pt=11, bold=True, underline=True)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    for i, instr in enumerate(instructions_list, 1):
        p = cell.add_paragraph()
        r = p.add_run(f'{i}. {instr}')
        set_run_font(r, size_pt=10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)


def build_questions_table(doc, sections):
    """Build the main 3-column table: Q.No | Question | Marks"""
    if not sections:
        return
    
    # Create main table
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Set column widths (total ~6.3 inches for A4 with 1" margins)
    # Col 0: Q No (0.5"), Col 1: Question (5.3"), Col 2: Marks (0.5")
    col_widths = [Inches(0.5), Inches(5.3), Inches(0.6)]
    
    border_thin = {'val': 'single', 'sz': 4, 'color': '000000'}
    border_none = {'val': 'none'}
    
    def add_table_row(q_text='', main_text='', marks_text='', 
                      q_bold=False, main_bold=False, marks_bold=False,
                      q_size=11, main_size=11, marks_size=11,
                      q_align=WD_ALIGN_PARAGRAPH.CENTER,
                      main_align=WD_ALIGN_PARAGRAPH.LEFT,
                      span_q_main=False, top_border=True, bottom_border=True,
                      bg_color=None, q_italic=False, main_italic=False,
                      main_indent=0, space_before=2, space_after=2,
                      marks_italic=False):
        row = table.add_row()
        cells = row.cells
        
        for ci, cell in enumerate(cells):
            cell.width = col_widths[ci]
            t_bdr = border_thin if top_border else border_none
            b_bdr = border_thin if bottom_border else border_none
            set_cell_border(cell, 
                           top=t_bdr, bottom=b_bdr,
                           left=border_thin, right=border_thin)
            if bg_color:
                set_cell_bg(cell, bg_color)
        
        # Q number cell
        p0 = cells[0].paragraphs[0]
        p0.alignment = q_align
        p0.paragraph_format.space_before = Pt(space_before)
        p0.paragraph_format.space_after = Pt(space_after)
        if q_text:
            r = p0.add_run(q_text)
            set_run_font(r, size_pt=q_size, bold=q_bold, italic=q_italic)
        
        # Main content cell
        p1 = cells[1].paragraphs[0]
        p1.alignment = main_align
        p1.paragraph_format.space_before = Pt(space_before)
        p1.paragraph_format.space_after = Pt(space_after)
        if main_indent:
            p1.paragraph_format.left_indent = Inches(main_indent)
        if main_text:
            r = p1.add_run(main_text)
            set_run_font(r, size_pt=main_size, bold=main_bold, italic=main_italic)
        
        # Marks cell
        p2 = cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(space_before)
        p2.paragraph_format.space_after = Pt(space_after)
        if marks_text:
            r = p2.add_run(str(marks_text))
            set_run_font(r, size_pt=marks_size, bold=marks_bold, italic=marks_italic)
        
        return cells
    
    # Header row
    hdr_cells = add_table_row('', 'Questions', 'Marks', 
                               main_bold=True, marks_bold=True, bg_color='E0E0E0',
                               main_align=WD_ALIGN_PARAGRAPH.CENTER, marks_size=11)
    
    q_counter = 1
    
    for section in sections:
        if not section.get('enabled', True):
            continue
        
        sec_letter = section.get('id', '')
        sec_name = section.get('name', '')
        sec_marks = section.get('totalMarks', '')
        sec_instructions = section.get('instructions', '')
        
        # Section header row
        section_text = f'SECTION {sec_letter}'
        if sec_name:
            section_text += f' ({sec_name})'
        
        add_table_row('', section_text, str(sec_marks) if sec_marks else '',
                     main_bold=True, marks_bold=True, bg_color='F5F5F5',
                     main_align=WD_ALIGN_PARAGRAPH.CENTER, main_size=12,
                     main_indent=0, space_before=4, space_after=4)
        
        # Section instructions
        if sec_instructions:
            add_table_row('', sec_instructions, '',
                         main_italic=True, main_size=10,
                         space_before=2, space_after=2)
        
        questions = section.get('questions', [])
        num_style = section.get('questionNumberStyle', '1, 2, 3...')
        
        for q_idx, question in enumerate(questions):
            q_type = question.get('type', 'sa')
            q_text = question.get('text', '')
            q_marks = question.get('marks', '')
            q_parts = question.get('parts', [])
            
            # Format question number
            q_num = format_q_number(q_counter, num_style)
            q_counter += 1
            
            # Handle passage/comprehension type
            if q_type in ['comprehension', 'unseen_passage']:
                add_table_row(q_num, q_text, str(q_marks) if q_marks else '',
                             q_bold=True, marks_bold=True,
                             main_bold=True, space_before=3, space_after=2)
                passage = question.get('passage', '')
                if passage:
                    add_table_row('', passage, '',
                                 main_italic=True, main_size=10,
                                 space_before=2, space_after=4)
            elif q_type in ['case_based']:
                add_table_row(q_num, q_text, str(q_marks) if q_marks else '',
                             q_bold=True, marks_bold=True,
                             main_bold=True, space_before=3, space_after=2)
                case_text = question.get('caseText', '')
                if case_text:
                    add_table_row('', case_text, '',
                                 main_italic=True, main_size=10,
                                 space_before=2, space_after=4)
            else:
                add_table_row(q_num, q_text, str(q_marks) if q_marks else '',
                             q_bold=True, marks_bold=True,
                             space_before=3, space_after=2)
            
            # MCQ options
            if q_type == 'mcq':
                options = question.get('options', [])
                opt_labels = ['a', 'b', 'c', 'd', 'e', 'f']
                for oi, opt in enumerate(options):
                    label = f'({opt_labels[oi]})' if oi < len(opt_labels) else f'({oi+1})'
                    add_table_row('', f'{label}  {opt}', '',
                                 main_size=11, main_indent=0.2,
                                 space_before=1, space_after=1)
            
            # Fill in blanks
            elif q_type == 'fill_blanks':
                blanks = question.get('blanks', [])
                for bi, blank in enumerate(blanks):
                    if blank:
                        add_table_row('', f'({bi+1})  {blank}', '',
                                     main_size=11, main_indent=0.2,
                                     space_before=1, space_after=1)
            
            # Match the following
            elif q_type == 'match':
                col_a = question.get('columnA', [])
                col_b = question.get('columnB', [])
                max_len = max(len(col_a), len(col_b))
                for mi in range(max_len):
                    a_item = f'{mi+1}. {col_a[mi]}' if mi < len(col_a) else ''
                    b_item = f'{chr(65+mi)}. {col_b[mi]}' if mi < len(col_b) else ''
                    combined = f'Column A: {a_item}     |     Column B: {b_item}'
                    add_table_row('', combined, '',
                                 main_size=10, main_indent=0.2,
                                 space_before=1, space_after=1)
            
            # Assertion-Reason
            elif q_type == 'assertion_reason':
                assertion = question.get('assertion', '')
                reason = question.get('reason', '')
                if assertion:
                    add_table_row('', f'Assertion (A): {assertion}', '',
                                 main_size=11, main_indent=0.1, space_before=2, space_after=1)
                if reason:
                    add_table_row('', f'Reason (R): {reason}', '',
                                 main_size=11, main_indent=0.1, space_before=1, space_after=2)
                ar_options = [
                    '(a) Both A and R are true and R is the correct explanation of A.',
                    '(b) Both A and R are true but R is not the correct explanation of A.',
                    '(c) A is true but R is false.',
                    '(d) A is false but R is true.'
                ]
                for opt in ar_options:
                    add_table_row('', opt, '',
                                 main_size=10, main_indent=0.2,
                                 space_before=1, space_after=1)
            
            # True/False
            elif q_type == 'true_false':
                add_table_row('', '(True / False)', '',
                             main_size=10, main_italic=True, main_indent=0.2,
                             space_before=1, space_after=2)
            
            # Parts
            for p_idx, part in enumerate(q_parts):
                part_label = chr(97 + p_idx)  # a, b, c...
                part_text = part.get('text', '')
                part_marks = part.get('marks', '')
                
                add_table_row('', f'({part_label})  {part_text}',
                             str(part_marks) if part_marks else '',
                             main_size=11, main_indent=0.2,
                             space_before=2, space_after=2)
                
                # Part sub-type handling
                part_type = part.get('type', '')
                if part_type == 'mcq':
                    opts = part.get('options', [])
                    opt_labels = ['i', 'ii', 'iii', 'iv']
                    for oi, opt in enumerate(opts):
                        lbl = f'({opt_labels[oi]})' if oi < len(opt_labels) else f'({oi+1})'
                        add_table_row('', f'{lbl}  {opt}', '',
                                     main_size=10, main_indent=0.4,
                                     space_before=1, space_after=1)
                
                # Subparts
                subparts = part.get('subparts', [])
                for sp_idx, subpart in enumerate(subparts):
                    sp_label = ['i', 'ii', 'iii', 'iv', 'v', 'vi'][sp_idx] if sp_idx < 6 else str(sp_idx+1)
                    sp_text = subpart.get('text', '') if isinstance(subpart, dict) else str(subpart)
                    sp_marks = subpart.get('marks', '') if isinstance(subpart, dict) else ''
                    add_table_row('', f'({sp_label})  {sp_text}',
                                 str(sp_marks) if sp_marks else '',
                                 main_size=10, main_indent=0.4,
                                 space_before=1, space_after=1)
            
            # Answer lines for SA/LA/VSA
            if q_type in ['sa', 'la', 'vsa', 'numerical'] and not q_parts:
                lines = question.get('answerLines', 0)
                if lines and int(lines) > 0:
                    for _ in range(int(lines)):
                        row = table.add_row()
                        for ci, cell in enumerate(row.cells):
                            cell.width = col_widths[ci]
                            set_cell_border(cell, 
                                          top=border_none, bottom={'val': 'single', 'sz': 2, 'color': 'AAAAAA'},
                                          left=border_thin, right=border_thin)
                        p = row.cells[1].paragraphs[0]
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(8)
        
        # Empty separator row after section
        add_table_row('', '', '', space_before=2, space_after=2,
                     top_border=False, bottom_border=False)


def format_q_number(counter, style):
    if style == 'Q1, Q2, Q3...' or style == 'Q1, Q2, Q3…':
        return f'Q{counter}.'
    elif style == 'i, ii, iii...' or style == 'i, ii, iii…':
        roman_map = {1:'i',2:'ii',3:'iii',4:'iv',5:'v',6:'vi',7:'vii',8:'viii',9:'ix',10:'x',
                     11:'xi',12:'xii',13:'xiii',14:'xiv',15:'xv',16:'xvi',17:'xvii',18:'xviii',19:'xix',20:'xx'}
        return roman_map.get(counter, str(counter)) + '.'
    elif style == '(a), (b), (c)...' or style == '(a), (b), (c)…':
        labels = [chr(96+i) for i in range(1, 27)]
        idx = (counter - 1) % 26
        return f'({labels[idx]})'
    else:
        return f'{counter}.'


def generate_exam_docx(data):
    doc = Document()
    
    # Page setup
    set_page_margins(doc)
    
    metadata = data.get('metadata', {})
    instructions = data.get('instructions', [])
    sections = data.get('sections', [])
    
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Build header table + instructions
    header_table, instr_cell = build_header_table(doc, metadata)
    add_instructions(instr_cell, instructions)
    
    # Footer with page numbers
    add_header_footer(doc, metadata)
    
    # Spacer paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    # Build questions table
    build_questions_table(doc, sections)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
